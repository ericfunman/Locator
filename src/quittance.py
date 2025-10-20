"""
Module de génération de quittances
"""

from docx import Document
from docx.shared import Pt, Inches
from datetime import datetime
from dateutil.relativedelta import relativedelta
import os
import shutil


def generer_quittance_simple(locataire, chambre, appartement, paiement, mois, annee):
    """
    Génère une quittance de loyer simple en format Word
    
    Args:
        locataire: Objet Locataire
        chambre: Objet Chambre
        appartement: Objet Appartement
        paiement: Objet Paiement
        mois: Numéro du mois (1-12)
        annee: Année
    
    Returns:
        Chemin du fichier généré
    """
    # Noms des mois en français
    mois_noms = [
        "Janvier", "Février", "Mars", "Avril", "Mai", "Juin",
        "Juillet", "Août", "Septembre", "Octobre", "Novembre", "Décembre"
    ]
    
    mois_nom = mois_noms[mois - 1] if 1 <= mois <= 12 else str(mois)
    
    # Créer un nouveau document
    doc = Document()
    
    # Titre
    title = doc.add_heading('QUITTANCE DE LOYER', 0)
    title.alignment = 1  # Centré
    
    # Espace
    doc.add_paragraph()
    
    # Informations du propriétaire (à personnaliser)
    doc.add_paragraph('Le bailleur :')
    doc.add_paragraph('[Votre nom]')
    doc.add_paragraph('[Votre adresse]')
    
    doc.add_paragraph()
    
    # Informations du locataire
    doc.add_paragraph(f'Certifie avoir reçu de :')
    doc.add_paragraph(f'{locataire.nom}')
    doc.add_paragraph(f'{appartement.adresse}')
    doc.add_paragraph(f'{appartement.code_postal} {appartement.ville}')
    
    doc.add_paragraph()
    
    # Période et montants
    doc.add_paragraph(f'La somme de : {paiement.montant:.2f} €')
    doc.add_paragraph(f'Pour la période du : {mois_nom} {annee}')
    
    doc.add_paragraph()
    doc.add_paragraph('Au titre de :')
    
    # Utiliser le montant du paiement pour avoir le montant historique correct
    montant_total = paiement.montant
    
    # Pour une quittance simple, on ne peut pas séparer loyer/charges
    # car on n'a que le montant total du paiement
    # On affiche juste le total
    
    # Tableau des détails
    table = doc.add_table(rows=2, cols=2)
    table.style = 'Light Grid Accent 1'
    
    # Total
    row = table.rows[0]
    row.cells[0].text = 'Loyer et charges'
    row.cells[1].text = f'{montant_total:.2f} €'
    
    # Total (répété pour clarté)
    row = table.rows[1]
    row.cells[0].text = 'TOTAL'
    row.cells[1].text = f'{montant_total:.2f} €'
    
    # Informations supplémentaires
    doc.add_paragraph()
    doc.add_paragraph()
    
    if paiement.date_paiement:
        doc.add_paragraph(f'Date du paiement : {paiement.date_paiement.strftime("%d/%m/%Y")}')
    
    if paiement.mode_paiement:
        doc.add_paragraph(f'Mode de paiement : {paiement.mode_paiement}')
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    # Signature
    doc.add_paragraph(f'Fait à [Ville], le {datetime.now().strftime("%d/%m/%Y")}')
    doc.add_paragraph()
    doc.add_paragraph('Signature du bailleur :')
    
    # Créer le nom de fichier
    safe_nom = locataire.nom.replace(' ', '_')
    filename = f"Quittance_{safe_nom}_{mois_nom}_{annee}.docx"
    temp_path = os.path.join("temp", filename)
    
    # Créer le dossier temp si nécessaire
    os.makedirs("temp", exist_ok=True)
    
    # Sauvegarder le document
    doc.save(temp_path)
    
    return temp_path


def generer_quittance_complete(locataire, bail, chambre, appartement, paiement, mois, annee, notes_supplementaires=""):
    """
    Génère une quittance complète à partir du template
    
    Args:
        locataire: Objet Locataire
        bail: Objet Bail
        chambre: Objet Chambre
        appartement: Objet Appartement
        paiement: Objet Paiement
        mois: Numéro du mois (1-12)
        annee: Année
        notes_supplementaires: Notes optionnelles
    
    Returns:
        Chemin du fichier généré
    """
    # Noms des mois en français
    mois_noms = [
        "Janvier", "Février", "Mars", "Avril", "Mai", "Juin",
        "Juillet", "Août", "Septembre", "Octobre", "Novembre", "Décembre"
    ]
    
    mois_nom = mois_noms[mois - 1] if 1 <= mois <= 12 else str(mois)
    
    # Chercher le template dans le dossier exemple
    exemple_dir = os.path.join("documents", "exemple")
    template_path = None
    
    # Chercher le premier fichier .docx dans le dossier exemple
    if os.path.exists(exemple_dir):
        for file in os.listdir(exemple_dir):
            if file.endswith('.docx') and not file.startswith('~'):
                template_path = os.path.join(exemple_dir, file)
                break
    
    if not template_path or not os.path.exists(template_path):
        # Si le template n'existe pas, créer une quittance simple
        return generer_quittance_simple(locataire, chambre, appartement, paiement, mois, annee)
    
    # Charger le template
    doc = Document(template_path)
    
    # Calculer les dates
    # Date début du mois en cours
    date_debut_mois = datetime(annee, mois, 1)
    
    # Date fin du mois en cours
    if mois == 12:
        date_fin_mois = datetime(annee + 1, 1, 1) - relativedelta(days=1)
        mois_suivant = 1
        annee_suivante = annee + 1
    else:
        date_fin_mois = datetime(annee, mois + 1, 1) - relativedelta(days=1)
        mois_suivant = mois + 1
        annee_suivante = annee
    
    # Date début du mois suivant
    date_debut_mois_suivant = datetime(annee_suivante, mois_suivant, 1)
    
    # Date fin du mois suivant
    if mois_suivant == 12:
        date_fin_mois_suivant = datetime(annee_suivante + 1, 1, 1) - relativedelta(days=1)
    else:
        date_fin_mois_suivant = datetime(annee_suivante, mois_suivant + 1, 1) - relativedelta(days=1)
    
    # Préparer les valeurs de remplacement
    nom_complet = locataire.nom
    adresse_complete = f"{appartement.adresse}, {appartement.code_postal} {appartement.ville}"
    
    # Utiliser le montant du paiement pour avoir le montant historique correct
    # et non le montant actuel du bail qui peut avoir changé
    montant_loyer = paiement.montant
    
    # Dictionnaire de remplacement avec les nouvelles variables
    replacements = {
        'DATE_DEBUT_MOIS': date_debut_mois.strftime('%d/%m/%Y'),
        'DATE_FIN_MOIS': date_fin_mois.strftime('%d/%m/%Y'),
        'DATE_DEBUT_SUIVANT_MOIS': date_debut_mois_suivant.strftime('%d/%m/%Y'),
        'DATE_FIN_SUIVANT_MOIS': date_fin_mois_suivant.strftime('%d/%m/%Y'),
        'ADRESSE_BIEN': adresse_complete,
        'NOM_LOCATAIRE': nom_complet,
        'MT_LOYER': f"{montant_loyer:.2f} €",
    }
    
    # Fonction pour remplacer dans un texte en gardant le formatage
    def remplacer_dans_runs(paragraph):
        """Remplace les variables dans les runs d'un paragraphe en gardant le formatage"""
        texte_complet = paragraph.text
        texte_modifie = texte_complet
        
        # Remplacer toutes les variables
        for key, value in replacements.items():
            texte_modifie = texte_modifie.replace(key, str(value))
        
        # Si le texte a changé, le remplacer
        if texte_modifie != texte_complet:
            # Garder le premier run et supprimer les autres
            for run in paragraph.runs[1:]:
                run.text = ''
            
            # Mettre le nouveau texte dans le premier run
            if paragraph.runs:
                paragraph.runs[0].text = texte_modifie
    
    # Remplacer dans tous les paragraphes
    for paragraph in doc.paragraphs:
        remplacer_dans_runs(paragraph)
    
    # Remplacer dans tous les tableaux
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    remplacer_dans_runs(paragraph)
    
    # Remplacer dans les en-têtes
    for section in doc.sections:
        if section.header:
            for paragraph in section.header.paragraphs:
                remplacer_dans_runs(paragraph)
        
        # Remplacer dans les pieds de page
        if section.footer:
            for paragraph in section.footer.paragraphs:
                remplacer_dans_runs(paragraph)
    
    # Créer le répertoire de destination
    safe_nom = locataire.nom.replace(' ', '_').replace("'", "").replace('-', '_')
    annee_str = str(annee)
    
    # Créer le chemin complet
    quittance_dir = os.path.join("documents", "locataires", safe_nom, annee_str, f"{mois:02d}")
    os.makedirs(quittance_dir, exist_ok=True)
    
    # Nom du fichier de sortie : quittance-nomlocataire-yyyymm.docx
    filename = f"quittance-{safe_nom.lower()}-{annee}{mois:02d}.docx"
    output_path = os.path.join(quittance_dir, filename)
    
    # Sauvegarder le document
    doc.save(output_path)
    
    return output_path


def nombre_en_lettres(nombre):
    """
    Convertit un nombre en lettres (version simplifiée pour les montants)
    """
    unites = ["", "un", "deux", "trois", "quatre", "cinq", "six", "sept", "huit", "neuf"]
    dizaines = ["", "dix", "vingt", "trente", "quarante", "cinquante", "soixante", "soixante-dix", "quatre-vingt", "quatre-vingt-dix"]
    
    nombre_entier = int(nombre)
    decimales = int(round((nombre - nombre_entier) * 100))
    
    if nombre_entier == 0:
        return "zéro"
    
    resultat = ""
    
    # Milliers
    milliers = nombre_entier // 1000
    reste = nombre_entier % 1000
    
    if milliers > 0:
        if milliers == 1:
            resultat = "mille "
        else:
            resultat = convertir_centaines(milliers) + " mille "
    
    # Centaines, dizaines, unités
    if reste > 0:
        resultat += convertir_centaines(reste)
    
    return resultat.strip()


def convertir_centaines(n):
    """Convertit un nombre de 0 à 999 en lettres"""
    unites = ["", "un", "deux", "trois", "quatre", "cinq", "six", "sept", "huit", "neuf"]
    dizaines = ["", "", "vingt", "trente", "quarante", "cinquante", "soixante", "soixante-dix", "quatre-vingt", "quatre-vingt-dix"]
    onze_a_seize = ["dix", "onze", "douze", "treize", "quatorze", "quinze", "seize"]
    
    if n == 0:
        return ""
    
    resultat = ""
    
    # Centaines
    centaines = n // 100
    reste = n % 100
    
    if centaines > 0:
        if centaines == 1:
            resultat = "cent "
        else:
            resultat = unites[centaines] + " cent "
    
    # Dizaines et unités
    if reste >= 10 and reste <= 16:
        resultat += onze_a_seize[reste - 10]
    elif reste >= 17 and reste <= 19:
        resultat += "dix-" + unites[reste - 10]
    elif reste >= 70 and reste <= 79:
        resultat += "soixante-" + onze_a_seize[reste - 70] if reste <= 76 else "soixante-dix-" + unites[reste - 70]
    elif reste >= 90 and reste <= 99:
        resultat += "quatre-vingt-" + onze_a_seize[reste - 90] if reste <= 96 else "quatre-vingt-dix-" + unites[reste - 90]
    else:
        dizaine = reste // 10
        unite = reste % 10
        
        if dizaine > 0:
            resultat += dizaines[dizaine]
            if unite > 0:
                if dizaine == 8:
                    resultat += "-" + unites[unite]
                elif unite == 1 and dizaine > 1:
                    resultat += " et un"
                else:
                    resultat += "-" + unites[unite]
        else:
            resultat += unites[unite]
    
    return resultat.strip()
